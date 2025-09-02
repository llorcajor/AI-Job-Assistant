import os
from dotenv import load_dotenv

# Load variables from the .env file
load_dotenv()

# apply.py
import json
import pandas as pd
import spacy
from docx import Document
from docx.shared import Inches
import google.generativeai as genai
from datetime import datetime
import os
import time
import random

# Selenium Imports
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from webdriver_manager.chrome import ChromeDriverManager
from bs4 import BeautifulSoup



# --- CONFIGURATION ---
API_KEY = os.getenv("API_KEY")
MASTER_RESUME_PATH = 'master_resume.json'
KNOWLEDGE_BASE_PATH = 'knowledge_base.json' # New
WRITING_STYLE_PATH = 'my_writing_style.txt' # New
TRACKING_FILE_PATH = 'seguimiento_aplicaciones.xlsx'
OUTPUT_BASE_FOLDER = 'applications'
NUM_BULLET_POINTS_PER_JOB = 4

# Configure the generative AI model
genai.configure(api_key=API_KEY)
model = genai.GenerativeModel('gemini-1.5-flash')


# --- HELPER FUNCTIONS (ALL MODULES) ---

def scrape_single_job_url(url):
    print(f"Scraping data from: {url}")
    options = Options()
    options.add_argument("user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36")
    service = Service(ChromeDriverManager().install())
    driver = webdriver.Chrome(service=service, options=options)
    driver.get(url)
    time.sleep(random.uniform(3, 5))
    soup = BeautifulSoup(driver.page_source, 'html.parser')
    
    job_data = {}
    company_url = None
    try:
        if "linkedin.com" in url:
            job_data['Puesto'] = soup.find('h1', class_='top-card-layout__title').get_text(strip=True)
            company_link_element = soup.find('a', class_='topcard__org-name-link')
            job_data['Empresa'] = company_link_element.get_text(strip=True)
            company_url = company_link_element['href']
            job_data['Ubicación'] = soup.find('span', class_='topcard__flavor topcard__flavor--bullet').get_text(strip=True)
            description_div = soup.find('div', class_='show-more-less-html__markup')
            job_data['Descripcion'] = description_div.get_text('\n', strip=True)
            job_data['Fuente'] = 'LinkedIn'
        else:
            print("Error: URL is not from a supported site.")
            driver.quit()
            return None, None
        
        job_data['Enlace'] = url
        driver.quit()
        return job_data, company_url
    except Exception as e:
        print(f"An error occurred while scraping the URL: {e}")
        driver.quit()
        return None, None

def scrape_about_page_text(company_url):
    """MODIFIED: Scrapes text from the specific URL you provide."""
    if not company_url: return None
    print(f"Researching company from: {company_url}")
    options = Options()
    options.add_argument("user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36")
    service = Service(ChromeDriverManager().install())
    driver = webdriver.Chrome(service=service, options=options)
    try:
        driver.get(company_url)
        time.sleep(random.uniform(2, 4))
        soup = BeautifulSoup(driver.page_source, 'html.parser')
        page_text = soup.find('body').get_text(' ', strip=True)
        driver.quit()
        return page_text
    except Exception as e:
        print(f"Could not research company website due to an error: {e}")
        driver.quit()
        return None
    
def generate_and_save_cover_letter(resume_data, job_data, folder_path, company_insights, knowledge_base, style_examples):
    """UPGRADED: Generates a cover letter using the knowledge base."""
    print("AI Writer is drafting the cover letter...")
    
    job_title = job_data['Puesto']
    company_name = job_data['Empresa']
    
    prompt = f"""
    Act as a professional career coach writing a cover letter for the {job_title} role at {company_name}.
    The tone must be professional and confident.
    Structure the letter using the "Problem-Solution" framework.

    First, identify the company's core need from the job description: "{job_data['Descripcion']}"
    Then, write the cover letter. Acknowledge their need and frame my experience as the direct solution.
    To make your points, you MUST select the most relevant accomplishment from the "Knowledge Base" provided below and weave its specific 'result' into the letter to provide a concrete, metric-driven example.
    Also, subtly reference the company's values, which are: "{company_insights}".

    Address the letter to "[Hiring Manager Name]".

    --- KNOWLEDGE BASE OF ACCOMPLISHMENTS ---
    {json.dumps(knowledge_base, indent=2)}
    --- END KNOWLEDGE BASE ---
    """
    
    initial_draft = model.generate_content(prompt).text
    final_draft = apply_writing_style(initial_draft, style_examples, "cover letter")

    filename = f"Cover_Letter_for_{company_name.replace(' ', '_')}.txt"
    save_path = os.path.join(folder_path, filename)
    with open(save_path, 'w') as f:
        f.write(final_draft)
    print(f"Successfully generated personalized cover letter: {save_path}")
    return save_path

# --- (Other functions like build_resume, setup_folder, log_application remain largely the same) ---
# ... (Remember to include the definitions for all other necessary helper 
    
def apply_writing_style(text_to_revise, style_examples, text_type="cover letter"):
    """UPGRADED: Uses your writing examples to revise the text."""
    print(f"Applying personal writing style to the {text_type}...")
    
    reviser_prompt = f"""
    You are a writing expert specializing in style transfer. Your goal is to revise an AI-generated text to match the personal writing style of the author.
    It should sound confident, professional, and human.

    First, carefully analyze the following examples of the author's writing to understand their tone, phrasing, and rhythm:
    --- STYLE EXAMPLES ---
    {style_examples}
    --- END STYLE EXAMPLES ---

    Now, using that style, revise the following AI-generated {text_type}.
    CRITICALLY, do not remove any key skills, metrics, or technical terms from the original text.

    --- AI DRAFT TO REVISE ---
    {text_to_revise}
    --- END AI DRAFT ---
    """
    
    try:
        response = model.generate_content(reviser_prompt)
        return response.text
    except Exception as e:
        print(f"Could not apply writing style due to an error: {e}")
        return text_to_revise

def get_company_insights(about_text):
    if not about_text: return "No company information available."
    print("AI Analyst is summarizing company info...")
    prompt = f"""You are a sharp business analyst. Read the following text from a company's website. Summarize the company's core mission, key values, and primary business goals in 3-4 bullet points. WEBSITE TEXT: --- {about_text[:4000]} ---"""
    response = model.generate_content(prompt)
    return response.text

def get_job_problem(job_description):
    print("AI Agent is identifying the core job problem...")
    prompt = f"""Analyze the following job description and identify the single biggest problem or need the company is trying to solve with this hire. State the problem in one concise sentence. JOB DESCRIPTION: --- {job_description} ---"""
    response = model.generate_content(prompt)
    return response.text.strip()

def build_resume_for_job(resume_data, job_keywords, company_name, folder_path, company_insights):
    """
    Builds the resume with an AI agent that dynamically rewrites each work experience
    to be tailored for the target job.
    """
    doc = Document()
    contact = resume_data['contact_info']
    doc.add_heading(contact['name'], level=0)
    doc.add_paragraph(f"{contact['email']} | {contact['phone']} | {contact['linkedin']}")
    
    # --- AI-Personalized Summary ---
    doc.add_heading('Professional Summary', level=1)
    summary_prompt = f"""
    Rewrite the following professional summary to be tailored for a job application at {company_name}.
    Subtly incorporate the company's values and mission, which are summarized here: "{company_insights}".
    Keep it concise and impactful. ORIGINAL SUMMARY: --- {resume_data['summary']} --- REVISED SUMMARY:
    """
    personalized_summary = model.generate_content(summary_prompt).text
    doc.add_paragraph(personalized_summary)
    
    # --- DYNAMIC WORK EXPERIENCE SECTION ---
    doc.add_heading('Work Experience', level=1)
    
    print("\n--- Rewriting work experiences to be tailored for the job ---")
    
    for job in resume_data['work_experience']:
        print(f"Tailoring experience from: {job['company']}...")
        
        # Prepare the prompt for the "Experience Rewriter" agent
        original_bullets = "\n- ".join([bullet['text'] for bullet in job['bullet_points']])
        
        experience_rewrite_prompt = f"""
        You are an expert resume writer. Your task is to rewrite a description of a past work experience to be highly relevant for a new job application.

        TARGET KEYWORDS FOR THE NEW JOB:
        ---
        {', '.join(job_keywords)}
        ---

        ORIGINAL WORK EXPERIENCE DATA:
        - Company: {job['company']}
        - Company Description: {job['company_description']}
        - My Accomplishments:
        - {original_bullets}
        ---

        INSTRUCTIONS:
        1.  First, write a single, concise paragraph (2-3 sentences) that describes the company. You MUST enhance this description by naturally weaving in some of the TARGET KEYWORDS.
        2.  Second, write 3-4 bullet points that describe my accomplishments. You MUST rewrite these bullet points to be more impactful and to integrate the TARGET KEYWORDS where they fit naturally. Focus on achievements and metrics.

        Provide only the rewritten paragraph and the bullet points.
        """
        
        # Call the AI to get the rewritten text
        rewritten_experience = model.generate_content(experience_rewrite_prompt).text
        
        # Add the rewritten experience to the document
        doc.add_heading(f"{job['title']} | {job['company']}", level=2)
        doc.add_paragraph(job['dates']).italic = True
        doc.add_paragraph(rewritten_experience)

    # --- (Add other sections like Skills and Education as before) ---
    doc.add_heading('Skills', level=1)
    # ... your skills logic here
    
    filename = f"CV_for_{company_name.replace(' ', '_')}.docx"
    save_path = os.path.join(folder_path, filename)
    doc.save(save_path)
    print(f"\nSuccessfully generated highly personalized resume: {save_path}")
    return save_path

def generate_and_save_cover_letter(resume_data, job_data, folder_path, company_insights, knowledge_base, style_examples):
    """UPGRADED: Generates a cover letter using the knowledge base and applies writing style."""
    print("AI Writer is drafting the cover letter...")
    
    job_title = job_data['Puesto']
    company_name = job_data['Empresa']
    
    # The "Writer" agent prompt
    writer_prompt = f"""
    Act as a professional career coach writing a cover letter for the {job_title} role at {company_name}.
    The tone must be professional and confident.
    Structure the letter using the "Problem-Solution" framework.

    First, identify the company's core need from the job description: "{job_data['Descripcion']}"
    Then, write the cover letter. Acknowledge their need and frame my experience as the direct solution.
    To make your points, you MUST select the most relevant accomplishment from the "Knowledge Base" provided below and weave its specific 'result' into the letter to provide a concrete, metric-driven example.
    Also, subtly reference the company's values, which are: "{company_insights}".

    Address the letter to "[Hiring Manager Name]".

    --- KNOWLEDGE BASE OF ACCOMPLISHMENTS ---
    {json.dumps(knowledge_base, indent=2)}
    --- END KNOWLEDGE BASE ---
    """
    
    initial_draft = model.generate_content(writer_prompt).text
    
    # The "Reviser" agent applies your personal writing style
    final_draft = apply_writing_style(initial_draft, style_examples, "cover letter")

    filename = f"Cover_Letter_for_{company_name.replace(' ', '_')}.txt"
    save_path = os.path.join(folder_path, filename)
    with open(save_path, 'w') as f:
        f.write(final_draft)
    print(f"Successfully generated personalized cover letter: {save_path}")
    return save_path, final_draft # Return both the path and the text

def setup_application_folder(job_data):
    """Creates a dedicated folder for the application files."""
    company_name = "".join(c for c in job_data['Empresa'] if c.isalnum() or c in (' ', '_')).rstrip()
    job_title = "".join(c for c in job_data['Puesto'] if c.isalnum() or c in (' ', '_')).rstrip()
    folder_name = f"{company_name.replace(' ', '_')}_{job_title.replace(' ', '_')}"
    folder_path = os.path.join(OUTPUT_BASE_FOLDER, folder_name)
    os.makedirs(folder_path, exist_ok=True)
    return folder_path

def log_application(job_data, folder_path):
    """Logs the application details to the Excel tracking file."""
    new_log_entry = {'Fecha de Aplicación': [datetime.now().strftime("%Y-%m-%d %H:%M")], 'Empresa': [job_data['Empresa']], 'Puesto': [job_data['Puesto']], 'Estado': ['Generated'], 'Ubicación': [job_data.get('Ubicación', 'N/A')], 'Fuente': [job_data.get('Fuente', 'N/A')], 'Enlace': [job_data['Enlace']], 'Carpeta de Archivos': [folder_path], 'Notas': ['']}
    new_log_df = pd.DataFrame(new_log_entry)
    if os.path.exists(TRACKING_FILE_PATH):
        tracking_df = pd.read_excel(TRACKING_FILE_PATH)
        updated_df = pd.concat([tracking_df, new_log_df], ignore_index=True)
    else:
        updated_df = new_log_df
    updated_df.to_excel(TRACKING_FILE_PATH, index=False)
    print(f"\nApplication successfully logged in '{TRACKING_FILE_PATH}'")

def extract_keywords_from_description(job_description, all_skills):
    """
    Uses spaCy to find which of our skills are mentioned in the job description.
    """
    # Load the pre-trained English language model
    nlp = spacy.load("en_core_web_sm")
    
    # Process the job description text to create a searchable object
    doc = nlp(job_description.lower())
    
    # Efficiently find and collect all skills that are present in the doc
    found_keywords = {skill for skill in all_skills if skill.lower() in doc.text}
            
    print(f"\nFound {len(found_keywords)} matching keywords: {', '.join(found_keywords)}")
    
    # Return the set of skills that were found
    return found_keywords

def run_application_process(job_url, company_url):
    """
    This function orchestrates the entire application generation process.
    It takes the URLs as input and returns the results.
    """
    # --- Step 1: Load All Data & Knowledge Bases ---
    try:
        with open(MASTER_RESUME_PATH, 'r') as f: resume_data = json.load(f)
        with open(KNOWLEDGE_BASE_PATH, 'r') as f: knowledge_base = json.load(f)
        with open(WRITING_STYLE_PATH, 'r') as f: style_examples = f.read()
    except FileNotFoundError as e:
        return {"success": False, "message": f"Error: Could not find a required file. {e}"}

    # --- Step 2: Scrape & Research Phase ---
    selected_job, scraped_company_url = scrape_single_job_url(job_url)
    
    if not selected_job:
        return {"success": False, "message": "Could not scrape the job URL. Please check the link."}
    
    about_text = scrape_about_page_text(company_url) # Uses the URL provided by the user
    company_insights = get_company_insights(about_text)


    all_my_skills = set(resume_data['skills']['technical_skills'] + resume_data['skills']['soft_skills'])
    matched_keywords = extract_keywords_from_description(selected_job['Descripcion'], all_my_skills)

    
    # --- Step 3: Writing & Organization Phase ---
    application_folder = setup_application_folder(selected_job)
    
    resume_filepath = build_resume_for_job(resume_data, matched_keywords, selected_job['Empresa'], application_folder, company_insights)
    
    cover_letter_filepath, cover_letter_text = generate_and_save_cover_letter(resume_data, selected_job, application_folder, company_insights, knowledge_base, style_examples)

    # --- Step 4: Logging Phase ---
    log_application(selected_job, application_folder)

    return {
        "success": True,
        "message": f"Successfully generated and logged application in folder: {application_folder}",
        "folder_path": application_folder,
        "cover_letter_text": cover_letter_text
    }